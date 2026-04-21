#!/usr/bin/env python3
"""Generate realistic OCR table examples for PDF/DOCX and produce a comparison report."""

from __future__ import annotations

import json
import os
import subprocess
import sys
import textwrap
from datetime import datetime
from pathlib import Path
from typing import Dict, Iterable, List

from docx import Document
from docx.shared import Inches
from PIL import Image, ImageDraw, ImageFont

ROOT = Path(__file__).resolve().parents[1]
OCR_DIR = ROOT / "test_results" / "ocr"
ASSETS_DIR = OCR_DIR / "assets"
REPORT_PATH = OCR_DIR / "ocr_table_examples_report.md"
RESULTS_JSON_PATH = OCR_DIR / "ocr_table_examples_results.json"

TABLE_IMAGE_PATH = ASSETS_DIR / "ocr_table_middle_source.png"
PDF_DEMO_PATH = ASSETS_DIR / "ocr_table_middle_pdf_demo.pdf"
PDF_PREVIEW_PATH = ASSETS_DIR / "ocr_table_middle_pdf_page2_preview.png"
DOCX_DEMO_PATH = ASSETS_DIR / "ocr_table_middle_docx_demo.docx"


def _load_font(size: int) -> ImageFont.ImageFont:
    candidates = [
        "DejaVuSans.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
    ]
    for candidate in candidates:
        try:
            return ImageFont.truetype(candidate, size=size)
        except Exception:
            continue
    return ImageFont.load_default()


def _draw_centered_text(
    draw: ImageDraw.ImageDraw,
    text: str,
    xy: tuple[int, int, int, int],
    font: ImageFont.ImageFont,
) -> None:
    x0, y0, x1, y1 = xy
    bbox = draw.textbbox((0, 0), text, font=font)
    width = max(1, bbox[2] - bbox[0])
    height = max(1, bbox[3] - bbox[1])
    x = x0 + (x1 - x0 - width) // 2
    y = y0 + (y1 - y0 - height) // 2
    draw.text((x, y), text, fill="black", font=font)


def _draw_left_text(
    draw: ImageDraw.ImageDraw,
    text: str,
    xy: tuple[int, int, int, int],
    font: ImageFont.ImageFont,
    left_padding: int = 16,
) -> None:
    x0, y0, x1, y1 = xy
    bbox = draw.textbbox((0, 0), text, font=font)
    height = max(1, bbox[3] - bbox[1])
    y = y0 + (y1 - y0 - height) // 2
    draw.text((x0 + left_padding, y), text, fill="black", font=font)


def create_table_image(path: Path) -> None:
    width, height = 1900, 1300
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)

    title_font = _load_font(58)
    subtitle_font = _load_font(34)
    header_font = _load_font(33)
    body_font = _load_font(31)

    draw.text((100, 60), "BON DE COMMANDE - PIECES TURBINE", fill="black", font=title_font)
    draw.text((100, 145), "Source: scan image integre dans un document", fill="black", font=subtitle_font)

    table_x = 100
    table_y = 250
    col_widths = [760, 220, 360, 360]
    row_height = 120

    headers = ["ITEM", "QTE", "PRIX EUR", "TOTAL EUR"]
    rows = [
        ["TURBINE A7", "2", "1250", "2500"],
        ["CAPTEUR X9", "5", "180", "900"],
        ["CABLE T1", "12", "25", "300"],
        ["SOUS TOTAL", "-", "-", "3700"],
        ["TVA 20", "-", "-", "740"],
        ["TOTAL TTC", "-", "-", "4440"],
    ]

    x_cursor = table_x
    for idx, header in enumerate(headers):
        cell = (x_cursor, table_y, x_cursor + col_widths[idx], table_y + row_height)
        draw.rectangle(cell, outline="black", width=5)
        _draw_centered_text(draw, header, cell, header_font)
        x_cursor += col_widths[idx]

    for row_idx, row_values in enumerate(rows, start=1):
        row_y0 = table_y + row_idx * row_height
        row_y1 = row_y0 + row_height
        x_cursor = table_x
        for col_idx, value in enumerate(row_values):
            cell = (x_cursor, row_y0, x_cursor + col_widths[col_idx], row_y1)
            draw.rectangle(cell, outline="black", width=4)
            if col_idx == 0:
                _draw_left_text(draw, value, cell, body_font)
            else:
                _draw_centered_text(draw, value, cell, body_font)
            x_cursor += col_widths[col_idx]

    draw.text((100, 1140), "REFERENCE DOSSIER: OCR-TABLE-DEMO-2026", fill="black", font=subtitle_font)
    path.parent.mkdir(parents=True, exist_ok=True)
    image.save(path)


def _draw_wrapped_text(
    draw: ImageDraw.ImageDraw,
    text: str,
    start_x: int,
    start_y: int,
    max_width: int,
    line_height: int,
    font: ImageFont.ImageFont,
) -> int:
    words = text.split()
    current_line: List[str] = []
    y = start_y

    for word in words:
        candidate = " ".join(current_line + [word])
        bbox = draw.textbbox((0, 0), candidate, font=font)
        width = bbox[2] - bbox[0]
        if width <= max_width:
            current_line.append(word)
            continue

        if current_line:
            draw.text((start_x, y), " ".join(current_line), fill="black", font=font)
            y += line_height
            current_line = [word]
        else:
            draw.text((start_x, y), word, fill="black", font=font)
            y += line_height

    if current_line:
        draw.text((start_x, y), " ".join(current_line), fill="black", font=font)
        y += line_height

    return y


def _build_scanned_page(title: str, body: str, table_image: Path | None = None) -> Image.Image:
    page = Image.new("RGB", (1800, 2400), "white")
    draw = ImageDraw.Draw(page)

    title_font = _load_font(60)
    text_font = _load_font(34)

    draw.text((120, 100), title, fill="black", font=title_font)
    y = _draw_wrapped_text(
        draw=draw,
        text=body,
        start_x=120,
        start_y=230,
        max_width=1550,
        line_height=58,
        font=text_font,
    )

    if table_image is not None:
        table = Image.open(table_image).convert("RGB")
        ratio = min(1500 / table.width, 1400 / table.height)
        resized = table.resize((int(table.width * ratio), int(table.height * ratio)))
        x = (page.width - resized.width) // 2
        y = max(y + 40, 520)
        page.paste(resized, (x, y))
        draw.text((120, y + resized.height + 35), "Tableau image non lisible sans OCR", fill="black", font=text_font)

    draw.text((120, 2260), "Archive maintenance 2026 - exemplaire scanne", fill="black", font=text_font)
    return page


def create_pdf_with_middle_table(pdf_path: Path, preview_path: Path, table_image: Path) -> None:
    page_1 = _build_scanned_page(
        title="DOSSIER MAINTENANCE - PAGE 1",
        body=(
            "Contexte: ce PDF est un scan. Le tableau de pieces n'est pas du texte natif et "
            "doit etre recuperable par OCR pour etre utile au parsing."
        ),
        table_image=None,
    )
    page_2 = _build_scanned_page(
        title="DOSSIER MAINTENANCE - PAGE 2 (TABLEAU)",
        body=(
            "Le tableau image ci-dessous est place au milieu du document. Sans OCR, le parser "
            "natif ne voit pas les lignes item/qte/prix/total."
        ),
        table_image=table_image,
    )
    page_3 = _build_scanned_page(
        title="DOSSIER MAINTENANCE - PAGE 3",
        body=(
            "Synthese: l'extraction doit conserver les informations cle des pieces facturees et "
            "les montants total TTC pour exploitation RAG et controle metier."
        ),
        table_image=None,
    )

    preview_path.parent.mkdir(parents=True, exist_ok=True)
    page_2.save(preview_path)

    pdf_path.parent.mkdir(parents=True, exist_ok=True)
    page_1.save(
        pdf_path,
        "PDF",
        resolution=220.0,
        save_all=True,
        append_images=[page_2, page_3],
    )


def create_docx_with_middle_table(docx_path: Path, table_image: Path) -> None:
    document = Document()
    document.add_heading("Compte rendu maintenance turbines - Avril 2026", level=1)
    document.add_paragraph(
        "Ce document DOCX contient du texte natif, mais aussi un tableau scanne en image "
        "au milieu de la page."
    )
    document.add_paragraph(
        "Objectif parsing: recuperer les colonnes ITEM, QTE, PRIX EUR et TOTAL EUR meme quand "
        "elles sont presentes uniquement dans une image."
    )
    document.add_paragraph("Tableau image integre:")
    document.add_picture(str(table_image), width=Inches(6.5))
    document.add_paragraph(
        "Commentaire apres tableau: verification metier attendue sur TOTAL TTC = 4440 EUR."
    )
    docx_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(docx_path)


def _parse_document(file_path: Path, extension: str, ocr_enabled: bool) -> Dict[str, object]:
    code = textwrap.dedent(
        """
        import json
        import pathlib
        import sys

        from parsing_core.extractors import extract_text_locally

        path = pathlib.Path(sys.argv[1])
        ext = sys.argv[2]
        result = extract_text_locally(path.read_bytes(), ext)
        print(json.dumps(result, ensure_ascii=False))
        """
    )

    env = os.environ.copy()
    env["OCR_ENABLED"] = "1" if ocr_enabled else "0"
    env["OCR_PDF_FALLBACK_ENABLED"] = "1" if ocr_enabled else "0"

    process = subprocess.run(
        [sys.executable, "-c", code, str(file_path), extension],
        cwd=str(ROOT),
        env=env,
        capture_output=True,
        text=True,
        check=False,
    )
    if process.returncode != 0:
        raise RuntimeError(
            f"Parsing failed for {file_path.name} (OCR {'ON' if ocr_enabled else 'OFF'}):\n"
            f"stdout={process.stdout}\n"
            f"stderr={process.stderr}"
        )

    output = process.stdout.strip().splitlines()
    if not output:
        raise RuntimeError(f"No parser output for {file_path.name}")
    return json.loads(output[-1])


def _extract_context_lines(text: str, keywords: Iterable[str], max_lines: int = 28) -> str:
    lines = [line.strip() for line in (text or "").splitlines() if line.strip()]
    if not lines:
        return "[empty]"

    upper_keywords = [keyword.upper() for keyword in keywords]
    hit_indices: List[int] = []

    for idx, line in enumerate(lines):
        upper_line = line.upper()
        if any(keyword in upper_line for keyword in upper_keywords):
            hit_indices.append(idx)

    if not hit_indices:
        return "\n".join(lines[:max_lines])

    selected = set()
    for hit in hit_indices:
        for idx in range(max(0, hit - 2), min(len(lines), hit + 3)):
            selected.add(idx)

    ordered = [lines[idx] for idx in sorted(selected)]
    return "\n".join(ordered[:max_lines])


def _compact_result(result: Dict[str, object], keywords: Iterable[str]) -> Dict[str, object]:
    text_value = str(result.get("text") or "")
    return {
        "parser_strategy": result.get("parser_strategy"),
        "parser_error": result.get("parser_error"),
        "ocr_attempted": result.get("ocr_attempted"),
        "ocr_used": result.get("ocr_used"),
        "ocr_pages": result.get("ocr_pages"),
        "ocr_supplement_pages": result.get("ocr_supplement_pages"),
        "ocr_engine_trace": result.get("ocr_engine_trace"),
        "text_length": len(text_value),
        "text_excerpt": _extract_context_lines(text_value, keywords),
    }


def _build_markdown_report(
    pdf_on: Dict[str, object],
    pdf_off: Dict[str, object],
    docx_on: Dict[str, object],
    docx_off: Dict[str, object],
) -> str:
    pdf_keywords = ["TURBINE", "CAPTEUR", "TOTAL TTC", "QTE", "PRIX"]
    docx_keywords = ["TURBINE", "CAPTEUR", "TOTAL TTC", "OCR embedded images", "Embedded image"]

    pdf_on_excerpt = _extract_context_lines(str(pdf_on.get("text") or ""), pdf_keywords)
    pdf_off_excerpt = _extract_context_lines(str(pdf_off.get("text") or ""), pdf_keywords)
    docx_on_excerpt = _extract_context_lines(str(docx_on.get("text") or ""), docx_keywords)
    docx_off_excerpt = _extract_context_lines(str(docx_off.get("text") or ""), docx_keywords)

    generated_at = datetime.utcnow().strftime("%Y-%m-%d %H:%M:%S UTC")

    return f"""# OCR Real Table Examples Report

- Generated at: {generated_at}
- Objective: provide concrete OCR proofs on table data hidden inside image content in PDF and DOCX.

## Example A - PDF with table image in the middle of the document

- Source file: assets/{PDF_DEMO_PATH.name}
- Preview (page 2):

![PDF page 2 with image table](assets/{PDF_PREVIEW_PATH.name})

### OCR ON (expected production mode)

- parser_strategy: {pdf_on.get('parser_strategy')}
- parser_error: {pdf_on.get('parser_error')}
- ocr_attempted: {pdf_on.get('ocr_attempted')}
- ocr_used: {pdf_on.get('ocr_used')}
- ocr_pages: {pdf_on.get('ocr_pages')}
- ocr_supplement_pages: {pdf_on.get('ocr_supplement_pages')}
- text_length: {len(str(pdf_on.get('text') or ''))}

```text
{pdf_on_excerpt}
```

### OCR OFF (control baseline)

- parser_strategy: {pdf_off.get('parser_strategy')}
- parser_error: {pdf_off.get('parser_error')}
- ocr_attempted: {pdf_off.get('ocr_attempted')}
- ocr_used: {pdf_off.get('ocr_used')}
- text_length: {len(str(pdf_off.get('text') or ''))}

```text
{pdf_off_excerpt}
```

## Example B - DOCX with embedded table image in the middle of content

- Source file: assets/{DOCX_DEMO_PATH.name}
- Embedded table image source:

![DOCX embedded image table source](assets/{TABLE_IMAGE_PATH.name})

### OCR ON (with embedded image OCR in DOCX parser)

- parser_strategy: {docx_on.get('parser_strategy')}
- parser_error: {docx_on.get('parser_error')}
- ocr_attempted: {docx_on.get('ocr_attempted')}
- ocr_used: {docx_on.get('ocr_used')}
- ocr_pages: {docx_on.get('ocr_pages')}
- ocr_supplement_pages: {docx_on.get('ocr_supplement_pages')}
- text_length: {len(str(docx_on.get('text') or ''))}

```text
{docx_on_excerpt}
```

### OCR OFF (control baseline)

- parser_strategy: {docx_off.get('parser_strategy')}
- parser_error: {docx_off.get('parser_error')}
- ocr_attempted: {docx_off.get('ocr_attempted')}
- ocr_used: {docx_off.get('ocr_used')}
- text_length: {len(str(docx_off.get('text') or ''))}

```text
{docx_off_excerpt}
```

## Key takeaway

These examples show the practical gain of OCR on non-native table content:
- PDF case: table text inside scanned page images is recoverable only when OCR is enabled.
- DOCX case: embedded image tables are now captured via OCR and appended to parsed output.
"""


def main() -> None:
    ASSETS_DIR.mkdir(parents=True, exist_ok=True)
    OCR_DIR.mkdir(parents=True, exist_ok=True)

    create_table_image(TABLE_IMAGE_PATH)
    create_pdf_with_middle_table(PDF_DEMO_PATH, PDF_PREVIEW_PATH, TABLE_IMAGE_PATH)
    create_docx_with_middle_table(DOCX_DEMO_PATH, TABLE_IMAGE_PATH)

    pdf_on = _parse_document(PDF_DEMO_PATH, "pdf", ocr_enabled=True)
    pdf_off = _parse_document(PDF_DEMO_PATH, "pdf", ocr_enabled=False)
    docx_on = _parse_document(DOCX_DEMO_PATH, "docx", ocr_enabled=True)
    docx_off = _parse_document(DOCX_DEMO_PATH, "docx", ocr_enabled=False)

    pdf_keywords = ["TURBINE", "CAPTEUR", "TOTAL TTC", "QTE", "PRIX"]
    docx_keywords = ["TURBINE", "CAPTEUR", "TOTAL TTC", "Embedded image", "OCR embedded images"]

    payload = {
        "generated_at": datetime.utcnow().isoformat() + "Z",
        "assets": {
            "table_image": str(TABLE_IMAGE_PATH.relative_to(ROOT)),
            "pdf_demo": str(PDF_DEMO_PATH.relative_to(ROOT)),
            "pdf_preview": str(PDF_PREVIEW_PATH.relative_to(ROOT)),
            "docx_demo": str(DOCX_DEMO_PATH.relative_to(ROOT)),
        },
        "comparisons": {
            "pdf_middle_table": {
                "ocr_on": _compact_result(pdf_on, pdf_keywords),
                "ocr_off": _compact_result(pdf_off, pdf_keywords),
            },
            "docx_middle_embedded_table": {
                "ocr_on": _compact_result(docx_on, docx_keywords),
                "ocr_off": _compact_result(docx_off, docx_keywords),
            },
        },
    }

    RESULTS_JSON_PATH.write_text(json.dumps(payload, indent=2, ensure_ascii=False), encoding="utf-8")
    REPORT_PATH.write_text(_build_markdown_report(pdf_on, pdf_off, docx_on, docx_off), encoding="utf-8")

    print("Generated files:")
    print(f"- {TABLE_IMAGE_PATH.relative_to(ROOT)}")
    print(f"- {PDF_DEMO_PATH.relative_to(ROOT)}")
    print(f"- {PDF_PREVIEW_PATH.relative_to(ROOT)}")
    print(f"- {DOCX_DEMO_PATH.relative_to(ROOT)}")
    print(f"- {RESULTS_JSON_PATH.relative_to(ROOT)}")
    print(f"- {REPORT_PATH.relative_to(ROOT)}")


if __name__ == "__main__":
    main()
